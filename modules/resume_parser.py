"""
Resume Parser — extracts structured text from PDF, DOCX, DOC, and TXT files.
Returns a clean text string + a basic metadata dict for the RAG context.
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────────
#  Extraction helpers
# ──────────────────────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    """Try PyMuPDF (fitz) first, then pdfplumber as fallback."""
    text = ""
    # Method 1: PyMuPDF (faster)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()
        if text.strip():
            return text
    except Exception as exc:
        logger.debug(f"PyMuPDF failed ({exc}), trying pdfplumber …")

    # Method 2: pdfplumber (better for tables)
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as exc:
        logger.error(f"PDF extraction failed: {exc}")
        return ""


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error(f"DOCX extraction failed: {exc}")
        return ""


def _extract_txt(path: str) -> str:
    try:
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        logger.error(f"TXT extraction failed: {exc}")
        return ""


# ──────────────────────────────────────────────────────────────────────────
#  Text cleaning
# ──────────────────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    # Collapse excessive whitespace but preserve paragraph breaks
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    return text.strip()


# ──────────────────────────────────────────────────────────────────────────
#  Metadata extraction (heuristic)
# ──────────────────────────────────────────────────────────────────────────

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.\w+")
_PHONE_RE = re.compile(r"\+?[\d\s\-().]{7,16}")
_URL_RE = re.compile(r"https?://\S+|linkedin\.com/\S+|github\.com/\S+", re.I)
_SKILL_KEYWORDS = {
    "python", "java", "javascript", "typescript", "react", "node", "sql",
    "aws", "azure", "gcp", "docker", "kubernetes", "machine learning",
    "deep learning", "tensorflow", "pytorch", "flask", "django", "fastapi",
    "spring", "git", "ci/cd", "devops", "agile", "scrum", "rest", "graphql",
    "html", "css", "mongodb", "postgresql", "redis", "kafka", "spark",
    "hadoop", "tableau", "power bi", "excel", "r", "scala", "go", "rust",
    "c++", "c#", "swift", "kotlin", "ruby", "php", "linux", "bash",
    "microservices", "system design", "data structures", "algorithms",
    "ai", "nlp", "computer vision", "llm", "rag", "langchain",
}


def _extract_metadata(text: str) -> dict:
    text_lower = text.lower()
    meta = {}

    # Email
    emails = _EMAIL_RE.findall(text)
    if emails:
        meta["email"] = emails[0]

    # Phone
    phones = _PHONE_RE.findall(text[:500])  # usually in header
    if phones:
        meta["phone"] = phones[0].strip()

    # URLs
    urls = _URL_RE.findall(text)
    meta["urls"] = list(set(urls[:5]))

    # Skills detection
    found_skills = []
    for skill in _SKILL_KEYWORDS:
        if skill in text_lower:
            found_skills.append(skill)
    meta["detected_skills"] = sorted(found_skills)

    # Rough experience estimate (look for "X years" patterns)
    exp_match = re.search(r"(\d+)\+?\s*years?\s*(of)?\s*(experience|exp)", text_lower)
    if exp_match:
        meta["years_experience"] = int(exp_match.group(1))

    # Education hints
    for edu_kw in ["bachelor", "master", "phd", "b.tech", "m.tech", "b.e", "m.e",
                   "bsc", "msc", "mba", "doctorate"]:
        if edu_kw in text_lower:
            meta["has_degree"] = True
            break

    return meta


# ──────────────────────────────────────────────────────────────────────────
#  Public API
# ──────────────────────────────────────────────────────────────────────────

def parse_resume(file_path: str) -> dict:
    """
    Parse a resume file and return:
    {
        "text":     <cleaned plain text>,
        "metadata": <heuristic metadata dict>,
        "word_count": <int>,
        "char_count": <int>,
        "success":  <bool>,
        "error":    <str | None>,
    }
    """
    path = Path(file_path)
    if not path.exists():
        return {"success": False, "error": f"File not found: {file_path}", "text": "", "metadata": {}}

    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            raw = _extract_pdf(str(path))
        elif suffix in (".docx",):
            raw = _extract_docx(str(path))
        elif suffix in (".doc",):
            # .doc requires antiword or LibreOffice; try docx fallback gracefully
            raw = _extract_docx(str(path))
            if not raw:
                raw = _extract_txt(str(path))
        elif suffix == ".txt":
            raw = _extract_txt(str(path))
        else:
            return {"success": False, "error": f"Unsupported file type: {suffix}", "text": "", "metadata": {}}

        if not raw or not raw.strip():
            return {"success": False, "error": "Could not extract text from file", "text": "", "metadata": {}}

        cleaned = _clean_text(raw)
        metadata = _extract_metadata(cleaned)

        return {
            "success": True,
            "error": None,
            "text": cleaned,
            "metadata": metadata,
            "word_count": len(cleaned.split()),
            "char_count": len(cleaned),
        }

    except Exception as exc:
        logger.exception(f"Resume parsing error for {file_path}: {exc}")
        return {"success": False, "error": str(exc), "text": "", "metadata": {}}


def parse_resume_text(text: str) -> dict:
    """Parse raw text (e.g. pasted resume content) instead of a file."""
    if not text or not text.strip():
        return {"success": False, "error": "Empty text provided", "text": "", "metadata": {}}
    cleaned = _clean_text(text)
    return {
        "success": True,
        "error": None,
        "text": cleaned,
        "metadata": _extract_metadata(cleaned),
        "word_count": len(cleaned.split()),
        "char_count": len(cleaned),
    }
